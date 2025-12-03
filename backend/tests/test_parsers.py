"""Tests for file parsers"""

import pytest
import tempfile
from pathlib import Path
from core.parsers import MarkdownParser, TextParser, DOCXParser
from docx import Document


def test_markdown_parser():
    """Test Markdown parsing"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write("# Heading 1\n\nContent here.\n\n## Heading 2\n")
        temp_path = f.name

    try:
        parser = MarkdownParser()
        result = parser.parse(temp_path)

        assert result.content
        assert len(result.headings) == 2
        assert 'Heading 1' in result.headings
        assert 'Heading 2' in result.headings
        assert 'encoding' in result.metadata
    finally:
        Path(temp_path).unlink()


def test_text_parser():
    """Test text file parsing"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write("Plain text content\nMultiple lines\n")
        temp_path = f.name

    try:
        parser = TextParser()
        result = parser.parse(temp_path)

        assert 'Plain text content' in result.content
        assert len(result.headings) == 0
        assert 'encoding' in result.metadata
    finally:
        Path(temp_path).unlink()


def test_docx_parser():
    """Test DOCX parsing"""
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_path = temp_file.name
    temp_file.close()

    try:
        # Create DOCX
        doc = Document()
        doc.add_heading('Test Heading', 1)
        doc.add_paragraph('Test paragraph')
        doc.save(temp_path)

        # Parse
        parser = DOCXParser()
        result = parser.parse(temp_path)

        assert 'Test Heading' in result.headings
        assert 'Test paragraph' in result.content
        assert isinstance(result.metadata, dict)
    finally:
        Path(temp_path).unlink()
