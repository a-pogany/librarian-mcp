"""File parsers for different document formats"""

from abc import ABC, abstractmethod
from typing import Dict, List
from pathlib import Path
import re
import chardet
from docx import Document as DOCXDocument


class ParseResult:
    """Result of parsing a document"""

    def __init__(self, content: str, headings: List[str], metadata: Dict):
        self.content = content
        self.headings = headings
        self.metadata = metadata

    def to_dict(self) -> Dict:
        return {
            'content': self.content,
            'headings': self.headings,
            'metadata': self.metadata
        }


class Parser(ABC):
    """Abstract parser interface"""

    @abstractmethod
    def parse(self, file_path: str) -> ParseResult:
        """
        Parse file and extract content

        Args:
            file_path: Path to file to parse

        Returns:
            ParseResult with content, headings, metadata
        """
        pass


class MarkdownParser(Parser):
    """Parse Markdown files"""

    def parse(self, file_path: str) -> ParseResult:
        """Parse Markdown file"""
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw = f.read()
            encoding = chardet.detect(raw)['encoding'] or 'utf-8'

        # Read file with detected encoding
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read()

        # Extract headings
        headings = self._extract_headings(content)

        return ParseResult(
            content=content,
            headings=headings,
            metadata={'encoding': encoding}
        )

    def _extract_headings(self, content: str) -> List[str]:
        """Extract markdown headings"""
        headings = []
        for line in content.split('\n'):
            if line.strip().startswith('#'):
                # Remove # symbols and clean
                heading = re.sub(r'^#+\s*', '', line).strip()
                if heading:
                    headings.append(heading)
        return headings


class TextParser(Parser):
    """Parse plain text files"""

    def parse(self, file_path: str) -> ParseResult:
        """Parse text file"""
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw = f.read()
            encoding = chardet.detect(raw)['encoding'] or 'utf-8'

        # Read file with detected encoding
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read()

        return ParseResult(
            content=content,
            headings=[],
            metadata={'encoding': encoding}
        )


class DOCXParser(Parser):
    """Parse DOCX files with enhanced metadata"""

    def parse(self, file_path: str) -> ParseResult:
        """Parse DOCX file with comprehensive metadata extraction"""
        doc = DOCXDocument(file_path)

        content_parts = []
        headings = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                headings.append(text)

            content_parts.append(text)

        # Extract tables
        for table in doc.tables:
            content_parts.append(self._format_table(table))

        # Extract comprehensive metadata
        metadata = self._extract_metadata(doc)

        return ParseResult(
            content='\n'.join(content_parts),
            headings=headings,
            metadata=metadata
        )

    def _format_table(self, table) -> str:
        """Format table as text"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))
        return '\n'.join(rows)

    def _extract_metadata(self, doc: DOCXDocument) -> Dict:
        """Extract comprehensive document metadata"""
        try:
            core_props = doc.core_properties

            # Count sections and estimate pages
            sections = self._extract_sections(doc)
            estimated_pages = self._estimate_pages(doc)
            heading_structure = self._extract_heading_structure(doc)

            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'revision': core_props.revision,
                'pages': estimated_pages,
                'sections': sections,
                'has_tables': len(doc.tables) > 0,
                'table_count': len(doc.tables),
                'headings': heading_structure
            }
        except Exception as e:
            return {'error': str(e)}

    def _extract_sections(self, doc: DOCXDocument) -> List[str]:
        """Extract section headings"""
        sections = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                sections.append(para.text.strip())
        return sections

    def _extract_heading_structure(self, doc: DOCXDocument) -> Dict[str, List[str]]:
        """Extract headings by level"""
        headings = {'h1': [], 'h2': [], 'h3': []}

        for para in doc.paragraphs:
            if para.style.name == 'Heading 1':
                headings['h1'].append(para.text.strip())
            elif para.style.name == 'Heading 2':
                headings['h2'].append(para.text.strip())
            elif para.style.name == 'Heading 3':
                headings['h3'].append(para.text.strip())

        return headings

    def _estimate_pages(self, doc: DOCXDocument) -> int:
        """
        Estimate page count (approximate)
        Assumptions: ~500 words per page
        """
        total_words = sum(len(para.text.split()) for para in doc.paragraphs)
        return max(1, total_words // 500)
